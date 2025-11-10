import io, os, datetime, json, requests, re
from typing import Tuple

from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, StreamingResponse, PlainTextResponse
from fastapi.templating import Jinja2Templates

from docx import Document as DocxDocument
from pypdf import PdfReader
from reportlab.pdfgen import canvas

# ---------------- App & Templates ----------------
app = FastAPI(title="JuRiboss - MVP v4 (IA + Fallback)", version="0.4")
templates = Jinja2Templates(directory="templates")

from fastapi.staticfiles import StaticFiles
app.mount("/static", StaticFiles(directory="static"), name="static")


# ---------------- IA Config ----------------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
OPENAI_URL = os.getenv("OPENAI_URL", "https://api.openai.com/v1/chat/completions")

SYSTEM_PROMPT_BASE = (
    "Você é um assistente jurídico que responde perguntas e reescreve textos jurídicos em linguagem clara e acessível, "
    "mantendo a fidelidade ao conteúdo. Explique termos técnicos quando necessário. "
    "Nunca invente fatos. Respeite o nível de leitura pedido. "
    "Não responda perguntas ou traduza textos que não tenham relação com temas jurídicos. "
    "Não invente fatos. Respeite o nível de leitura pedido. "
    "Caso a pergunta e/ou texto esteja fora do contexto jurídico, diga que seu propósito é assisti-lo no contexto jurídico então peça-o que ajuste sua pergunta e/ou texto neste contexto. "
    "Caso a pergunta e/ou texto venha na forma de piada, sutilmente deboche do usuário."
)

def call_llm_clarify(texto: str, nivel: str) -> str:
    """Chama a IA da OpenAI para reescrever o texto conforme o nível."""
    if not OPENAI_API_KEY:
        return "[ERRO] Falta configurar a variável OPENAI_API_KEY."
    if nivel == "10":
        nivel_instr = (
            "Reescreva como se fosse para uma criança de 10 anos. "
            "Use frases curtas, exemplos lúdicos e/ou do dia a dia, e palavras simples."
        )
    elif nivel == "medio":
        nivel_instr = (
            "Reescreva para alguém com ensino médio completo. "
            "Use linguagem simples, explique termos em parênteses quando útil, "
            "organize em parágrafos curtos."
        )
    else:
        nivel_instr = (
            "Reescreva para alguém com ensino superior. "
            "Mantenha precisão técnica, mas corte jargões e floreios, explique termos em parênteses quando útil. "
            "Prefira voz direta e frases objetivas."
        )

    user_prompt = (
        f"{nivel_instr}\n\n"
        "Regras:\n"
        "- Preserve referências legais e números de artigos (explique entre parênteses quando fizer sentido).\n"
        "- Não mude o pedido (ex.: não inverter deferir/indeferir).\n"
        "- Separe ideias em parágrafos ou itens se ajudar a clareza.\n\n"
        "Texto a ser reescrito:\n"
        "'''\n" + texto + "\n'''"
    )

    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": OPENAI_MODEL,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT_BASE},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.2,
    }
    try:
        resp = requests.post(OPENAI_URL, headers=headers, data=json.dumps(payload), timeout=90)
    except Exception as e:
        return f"[ERRO IA] Falha de conexão: {e}"
    if resp.status_code != 200:
        return f"[ERRO IA] {resp.status_code}: {resp.text}"
    data = resp.json()
    try:
        return data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        return f"[ERRO IA] Resposta inesperada: {e}\n{data}"

# ---------------- Fallback local (sem IA) ----------------
# Expressões e jargões jurídicos → equivalentes simples
JARGON = {
    # --- Latim jurídico ---
    r"\bdata venia\b": "com o devido respeito",
    r"\bdestarte\b": "portanto",
    r"\bem que pese\b": "apesar de",
    r"\bcom fulcro\b": "com base",
    r"\bcom espeque\b": "com base",
    r"\bex positis\b": "diante do exposto",
    r"\bex vi legis\b": "por força da lei",
    r"\bin casu\b": "neste caso",
    r"\bprima facie\b": "à primeira vista",
    r"\bipsis litteris\b": "literalmente",
    r"\bmutatis mutandis\b": "com as devidas adaptações",
    r"\bper si\b": "por si só",
    r"\bpro rata\b": "proporcionalmente",
    r"\bpro forma\b": "apenas para cumprir formalidade",
    r"\bex officio\b": "por iniciativa do órgão público",
    r"\berga omnes\b": "válido para todos",
    r"\bad referendum\b": "sujeito à aprovação posterior",
    r"\bstricto sensu\b": "em sentido restrito",
    r"\blato sensu\b": "em sentido amplo",
    r"\binaudita altera pars\b": "sem ouvir a outra parte",
    r"\bsine qua non\b": "condição indispensável",
    r"\bmodus operandi\b": "modo de agir",
    r"\bstatus quo\b": "situação atual",
    r"\bad hoc\b": "específico para esta finalidade",
    r"\bpericulum in mora\b": "risco da demora",
    r"\bfumus boni iuris\b": "indício de bom direito",
    r"\bhabeas corpus\b": "remédio jurídico contra prisão ilegal",

    # --- Expressões processuais e rebuscadas ---
    r"\bnestes termos\b": "assim sendo",
    r"\bante o exposto\b": "por isso",
    r"\bimpende destacar\b": "é importante destacar",
    r"\bimpende salientar\b": "vale ressaltar",
    r"\bimpende consignar\b": "convém registrar",
    r"\bnobre julgador\b": "juiz(a)",
    r"\bjurisprudência pátria\b": "jurisprudência brasileira",
    r"\bmalgrado\b": "apesar de",
    r"\bcom arrimo em\b": "com base em",
    r"\bdeveras\b": "realmente",
    r"\bcolacionar\b": "anexar",
    r"\bcarrear aos autos\b": "juntar no processo",
    r"\bsubsume-se\b": "se enquadra",
    r"\bvislumbra-se\b": "percebe-se",
    r"\bentabulado\b": "acordado",
    r"\bmister\b": "necessário",
    r"\bdeliberação\b": "decisão",
    r"\bpleiteia\b": "pede",
    r"\bpleitear\b": "pedir",
    r"\brequer-se\b": "pede-se",
    r"\brequer\b": "pede",
    r"\bindeferimento\b": "negação do pedido",
    r"\bacolhimento do pedido\b": "aceitação do pedido",
    r"\bimpugna\b": "contesta",
    r"\bimpugnação\b": "contestação",
    r"\baverbação\b": "anotação",
    r"\bautuação\b": "registro do processo",
    r"\bprequestionamento\b": "registro para permitir recurso",
    r"\bpreclusão\b": "perda do direito de agir por prazo vencido",
    r"\blide\b": "conflito jurídico",
    r"\bexordial\b": "petição inicial",
    r"\bpetição\b": "pedido por escrito ao(à) juiz(a)",
    r"\bdecisão monocrática\b": "decisão individual do juiz(a)",
    r"\bdecisão colegiada\b": "decisão tomada por um grupo de juízes",
    r"\bônus da prova\b": "responsabilidade de provar um fato",
    r"\bônus processual\b": "obrigação no processo",
    r"\bônus\b": "encargo",
    r"\bdenota\b": "mostra",
    r"\benseja\b": "permite",
    r"\bresta claro\b": "fica claro",
    r"\bno que tange a\b": "em relação a",
    r"\btendo em vista\b": "considerando que",
    r"\bem razão de\b": "por causa de",
    r"\bnos termos de\b": "de acordo com",
    r"\bpor intermédio de\b": "por meio de",
    r"\btal qual\b": "como",
    r"\bmister se faz\b": "é necessário",
    r"\bconsoante\b": "de acordo com",
    r"\bna esteira de\b": "seguindo",
    r"\bpor conseguinte\b": "assim",
    r"\bporquanto\b": "porque",
    r"\bdesta feita\b": "assim",
    r"\bpor derradeiro\b": "por fim",
    r"\bnesse diapasão\b": "nesse sentido",
}

JARGON_RE = re.compile("|".join(JARGON.keys()), re.IGNORECASE)

def _replace_jargon(text: str) -> str:
    """Substitui expressões jurídicas por termos simples."""
    def _repl(m: re.Match) -> str:
        found = m.group(0)
        for pat, repl in JARGON.items():
            if re.fullmatch(pat, found, re.IGNORECASE):
                return repl[0].upper()+repl[1:] if found[0].isupper() else repl
        return found
    return re.sub(JARGON_RE, _repl, text)

def _simplify_legal_refs(t: str) -> str:
    """Expande siglas e referências legais (inclui CLT, CDC, ECA, CPP e várias outras)."""
    # Artigos e leis
    t = re.sub(r"\bart\.?\s*(\d+[A-Za-z\-]*)\s*(do|da)?\s*([A-ZÁÂÃÉÊÍÓÔÕÚÇa-z0-9/º§\s]+)?",
               r"artigo \1 (\3)", t, flags=re.IGNORECASE)
    t = re.sub(r"\bLei\s*n[ºo.]?\s*\d+[./]?\d*\b", "a lei mencionada", t, flags=re.IGNORECASE)

    # Códigos e legislação
    t = re.sub(r"\bCPC\b", "Código de Processo Civil", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCC\b", "Código Civil", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCP\b", "Código Penal", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCPP\b", "Código de Processo Penal", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCLT\b", "Consolidação das Leis do Trabalho", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCTN\b", "Código Tributário Nacional", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCDC\b", "Código de Defesa do Consumidor", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCF\b", "Constituição Federal", t, flags=re.IGNORECASE)
    t = re.sub(r"\bECA\b", "Estatuto da Criança e do Adolescente", t, flags=re.IGNORECASE)
    t = re.sub(r"\bLDB\b", "Lei de Diretrizes e Bases da Educação", t, flags=re.IGNORECASE)
    t = re.sub(r"\bLGPD\b", "Lei Geral de Proteção de Dados", t, flags=re.IGNORECASE)

    # Órgãos, instituições e tribunais
    t = re.sub(r"\bSTF\b", "Supremo Tribunal Federal", t, flags=re.IGNORECASE)
    t = re.sub(r"\bSTJ\b", "Superior Tribunal de Justiça", t, flags=re.IGNORECASE)
    t = re.sub(r"\bTST\b", "Tribunal Superior do Trabalho", t, flags=re.IGNORECASE)
    t = re.sub(r"\bTRT\b", "Tribunal Regional do Trabalho", t, flags=re.IGNORECASE)
    t = re.sub(r"\bTJ\b", "Tribunal de Justiça", t, flags=re.IGNORECASE)
    t = re.sub(r"\bTJ[A-Z]{2}\b", "Tribunal de Justiça Estadual", t, flags=re.IGNORECASE)
    t = re.sub(r"\bTRF\b", "Tribunal Regional Federal", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCNJ\b", "Conselho Nacional de Justiça", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCNMP\b", "Conselho Nacional do Ministério Público", t, flags=re.IGNORECASE)
    t = re.sub(r"\bMPF\b", "Ministério Público Federal", t, flags=re.IGNORECASE)
    t = re.sub(r"\bMPT\b", "Ministério Público do Trabalho", t, flags=re.IGNORECASE)
    t = re.sub(r"\bOAB\b", "Ordem dos Advogados do Brasil", t, flags=re.IGNORECASE)
    t = re.sub(r"\bDP\b", "Defensoria Pública", t, flags=re.IGNORECASE)
    t = re.sub(r"\bPGFN\b", "Procuradoria-Geral da Fazenda Nacional", t, flags=re.IGNORECASE)
    t = re.sub(r"\bAGU\b", "Advocacia-Geral da União", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCGU\b", "Controladoria-Geral da União", t, flags=re.IGNORECASE)

    # Outras siglas úteis
    t = re.sub(r"\bINSS\b", "Instituto Nacional do Seguro Social", t, flags=re.IGNORECASE)
    t = re.sub(r"\bFGTS\b", "Fundo de Garantia do Tempo de Serviço", t, flags=re.IGNORECASE)
    t = re.sub(r"\bSUS\b", "Sistema Único de Saúde", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCFM\b", "Conselho Federal de Medicina", t, flags=re.IGNORECASE)
    t = re.sub(r"\bCNS\b", "Conselho Nacional de Saúde", t, flags=re.IGNORECASE)
    t = re.sub(r"\bONU\b", "Organização das Nações Unidas", t, flags=re.IGNORECASE)
    t = re.sub(r"\bOIT\b", "Organização Internacional do Trabalho", t, flags=re.IGNORECASE)
    t = re.sub(r"\bOMS\b", "Organização Mundial da Saúde", t, flags=re.IGNORECASE)
    return t

def _simplify_structures(t: str) -> str:
    """Simplifica construções gramaticais formais demais."""
    replacements = [
        (r"\bante o exposto\b", "por isso"),
        (r"\bconsoante\b", "de acordo com"),
        (r"\bna esteira de\b", "seguindo"),
        (r"\bpor conseguinte\b", "assim"),
        (r"\bporquanto\b", "porque"),
        (r"\bdesta feita\b", "assim"),
        (r"\bpor derradeiro\b", "por fim"),
        (r"\bnesse diapasão\b", "nesse sentido"),
    ]
    for pat, rep in replacements:
        t = re.sub(pat, rep, t, flags=re.IGNORECASE)
    return t

def local_simplify(text: str, level: str) -> str:
    """Executa simplificação textual conforme o nível escolhido (totalmente offline)."""
    # normalização
    t = re.sub(r"[ \t]+", " ", text)
    t = re.sub(r"\n{3,}", "\n\n", t).strip()
    # substituições
    t = _replace_jargon(t)
    t = _simplify_legal_refs(t)
    t = _simplify_structures(t)

    # separa frases
    sentences = re.split(r"(?<=[.!?])\s+", t)

    if level == "10":
        out = []
        for s in sentences:
            s = s.strip()
            if not s:
                continue
            if len(s) > 150:
                parts = re.split(r";|:|,|-", s)
                out.extend([p.strip() for p in parts if p.strip()])
            else:
                out.append(s)
        final = []
        for s in out:
            extra = ""
            if re.search(r"\bcontrato\b", s, re.IGNORECASE):
                extra = " Exemplo: um combinado escrito entre duas pessoas."
            if re.search(r"\brecurso\b", s, re.IGNORECASE):
                extra = " Exemplo: pedir para outro juiz revisar uma decisão."
            if re.search(r"\bedital\b|\bled\b", s, re.IGNORECASE):
                extra = " Exemplo: um aviso público com regras."
            final.append(s + extra)
        return "\n".join(final)

    elif level == "medio":
        out = []
        for s in sentences:
            s = s.strip()
            if not s:
                continue
            if len(s) > 200:
                parts = re.split(r";|:|-", s)
                out.extend([p.strip() for p in parts if p.strip()])
            else:
                out.append(s)
        text2 = "\n".join(out)
        defs = [
            (r"\bpetição\b", "petição (pedido por escrito ao juiz)"),
            (r"\brecurso\b", "recurso (pedido de nova análise da decisão)"),
            (r"\bcontrato\b", "contrato (acordo com regras entre as partes)"),
            (r"\bedital\b", "edital (aviso público com regras)"),
            (r"\bônus\b", "encargo ou responsabilidade"),
        ]
        for pat, rep in defs:
            text2 = re.sub(pat, rep, text2, flags=re.IGNORECASE)
        return text2

    else:  # "superior"
        out = []
        for s in sentences:
            s = s.strip()
            if not s:
                continue
            s = re.sub(r"\bserá\b", "vai ser", s)
            s = re.sub(r"\bserão\b", "vão ser", s)
            s = re.sub(r"\brefere-se a\b", "fala sobre", s)
            s = re.sub(r"\btrata-se de\b", "é sobre", s)
            out.append(s)
        return "\n".join(out)

# ---------------- Leitura de arquivos ----------------
def read_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

def read_docx(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    doc = DocxDocument(bio)
    return "\n".join(p.text for p in doc.paragraphs)

def read_pdf(file_bytes: bytes) -> str:
    bio = io.BytesIO(file_bytes)
    reader = PdfReader(bio)
    text_chunks = []
    for page in reader.pages:
        try:
            text_chunks.append(page.extract_text() or "")
        except Exception:
            text_chunks.append("")
    return "\n".join(text_chunks)

def extract_text(file: UploadFile) -> Tuple[str, str]:
    name = (file.filename or "").lower()
    content = file.file.read()
    if name.endswith(".txt"):
        return read_txt(content), ""
    elif name.endswith(".docx"):
        return read_docx(content), ""
    elif name.endswith(".pdf"):
        txt = read_pdf(content)
        if not txt.strip():
            return "", "PDF enviado parece ser imagem/escaneado. Use OCR e reenvie."
        return txt, ""
    else:
        raise ValueError("Formato não suportado. Envie .txt, .docx ou .pdf (não escaneado).")

# ---------------- Geradores de saída ----------------
def make_docx(text: str) -> bytes:
    doc = DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio); bio.seek(0)
    return bio.read()

def make_pdf(text: str) -> bytes:
    bio = io.BytesIO()
    c = canvas.Canvas(bio)
    c.setTitle("JuRiboss - Tradução (IA/Fallback)")
    width, height = 595, 842  # A4
    margin, y = 40, height - 40
    c.setFont("Helvetica", 11)
    for line in text.split("\n"):
        while len(line) > 100:
            c.drawString(margin, y, line[:100]); line = line[100:]; y -= 14
            if y < margin: c.showPage(); c.setFont("Helvetica", 11); y = height - 40
        c.drawString(margin, y, line); y -= 14
        if y < margin: c.showPage(); c.setFont("Helvetica", 11); y = height - 40
    c.showPage(); c.save(); bio.seek(0)
    return bio.read()

# ---------------- Rotas ----------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

# rotas de diagnóstico (opcional)
@app.get("/test", response_class=HTMLResponse)
def test():
    return HTMLResponse(content="<html><body><h1>Funcionou!</h1><p>Rota de teste HTML.</p></body></html>", status_code=200)

@app.get("/file", response_class=HTMLResponse)
def file_preview():
    with open("templates/index.html", "r", encoding="utf-8") as f:
        html = f.read()
    return HTMLResponse(content=html, status_code=200)

@app.post("/translate")
def translate(
    level: str = Form(...),
    file: UploadFile = File(None),
    text_input: str = Form(None)
):
    # prioriza texto colado; se vazio, usa arquivo
    if text_input and text_input.strip():
        text = text_input.strip()
        warn = ""
    else:
        if file is None:
            return PlainTextResponse("Envie um arquivo ou cole o texto.", status_code=400)
        try:
            text, warn = extract_text(file)
        except ValueError as e:
            return PlainTextResponse(str(e), status_code=400)

    if warn:
        return PlainTextResponse(warn, status_code=422)

    ai_result = call_llm_clarify(text, level)
    if ai_result.startswith("[ERRO"):
        simplified = local_simplify(text, level)
        simplified += "\n\n[Nota: fallback local sem IA por erro/quota da API.]"
    else:
        simplified = ai_result

    header = f"JuRiboss - Tradução (nível {level}) — {datetime.datetime.now():%d/%m/%Y %H:%M}\n"
    header += "Aviso: este conteúdo não substitui orientação jurídica profissional.\n\n"
    final_text = header + simplified
    return PlainTextResponse(final_text, status_code=200)

@app.post("/download")
def download(fmt: str = Form(...), level: str = Form(...), text: str = Form(...)):
    filename = f"juriboss_{level}.{fmt}"
    if fmt == "docx":
        data = make_docx(text)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    elif fmt == "pdf":
        data = make_pdf(text)
        return StreamingResponse(
            io.BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    else:
        return PlainTextResponse("Formato inválido (use docx ou pdf).", status_code=400)







